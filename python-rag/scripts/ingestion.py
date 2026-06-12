"""
PHASE 1 + PHASE 2 : INGESTION COMPLETE
- Phase 1 : Lecture, nettoyage, decoupage
- Phase 2 : Embeddings + Stockage dans Vector DB

Objectifs qualite ajoutes:
- Lire plus de formats metier: pdf, docx, txt, xlsx, csv, html, png, jrxml
- Masquer les colonnes sensibles (ex: password/mot de passe) pour csv/xlsx
- Produire un audit d'extraction par document
- Produire un rapport de couverture (indexed_files / total_files) par categorie et extension
"""

import csv
import hashlib
import io
import os
import random
import re
import shutil
import sys
import zipfile
import xml.etree.ElementTree as ET
from datetime import datetime
from html import unescape
from pathlib import Path

# Ajouter le repertoire parent au path
sys.path.insert(0, str(Path(__file__).parent.parent))

# Charger les variables d'environnement depuis .env (GOOGLE_API_KEY, etc.)
try:
    from dotenv import load_dotenv
    load_dotenv(dotenv_path=Path(__file__).parent.parent / ".env")
except ImportError:
    pass

try:
    import chromadb
except ImportError:
    chromadb = None

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

import PyPDF2
from docx import Document
from tqdm import tqdm

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    import openpyxl
except ImportError:
    openpyxl = None

try:
    from bs4 import BeautifulSoup
except ImportError:
    BeautifulSoup = None

try:
    from PIL import Image
except ImportError:
    Image = None

try:
    import pytesseract
except ImportError:
    pytesseract = None

try:
    import easyocr
except ImportError:
    easyocr = None


_EASYOCR_READER = None


# ===== CONFIGURATION =====
DOCUMENTS_PATH = "./src/document"
VECTOR_DB_PATH = "./vector_db"

# 800 chars = bon compromis pour garder les exigences AFD entières (ID + Description + Règles)
CHUNK_SIZE = 800
CHUNK_OVERLAP = 100
LOW_TEXT_THRESHOLD = 50
RUN_PHASE2 = os.getenv("RUN_PHASE2", "1").strip().lower() not in ("0", "false", "no")
PRINT_STEP_RESULTS = os.getenv("PRINT_STEP_RESULTS", "1").strip().lower() not in (
    "0",
    "false",
    "no",
)
STEP_PREVIEW_CHARS = int(os.getenv("STEP_PREVIEW_CHARS", "500"))
STEP_CHUNKS_PREVIEW = int(os.getenv("STEP_CHUNKS_PREVIEW", "2"))

# Extraction tables PDF
PDF_EXTRACT_TABLES = os.getenv("PDF_EXTRACT_TABLES", "1").strip().lower() not in (
    "0",
    "false",
    "no",
)

# Embeddings (paraphrase-multilingual-MiniLM-L12-v2 : pas de préfixe requis)
EMBEDDING_MODEL_NAME = os.getenv(
    "EMBEDDING_MODEL_NAME", "paraphrase-multilingual-MiniLM-L12-v2"
).strip()

SUPPORTED_EXTENSIONS = (
    ".pdf",
    ".docx",
    ".txt",
    ".xlsx",
    ".csv",
    ".html",
    ".htm",
    ".png",
    ".jpg",
    ".jpeg",
    ".jrxml",
    ".xml",
)

TEXT_FILE_ENCODINGS = ["utf-8", "utf-8-sig", "cp1252", "latin-1"]

SENSITIVE_COLUMN_KEYWORDS = (
    "password",
    "motdepasse",
    "passwd",
    "mdp",
    "secret",
    "token",
    "apikey",
    "accesskey",
)
SENSITIVE_VALUE_MASK = "[REDACTED]"

PASSWORD_ASSIGN_PATTERN = re.compile(
    r"(?i)\b(mot\s*de\s*passe|password|passwd|mdp)\b(\s*[:=]\s*)([^\s,;|]+)"
)
SECRET_ASSIGN_PATTERN = re.compile(
    r"(?i)\b(api[_ -]?key|access[_ -]?key|token|secret|bearer)\b(\s*[:=]\s*)([^\s,;|]+)"
)
EMAIL_PATTERN = re.compile(r"(?i)\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b")
PHONE_PATTERN = re.compile(r"(?<!\d)(?:\+?\d[\d .()\-]{7,}\d)(?!\d)")
JWT_PATTERN = re.compile(r"\beyJ[A-Za-z0-9_-]{8,}\.[A-Za-z0-9._-]{8,}\.[A-Za-z0-9._-]{8,}\b")

AUDIT_OUTPUT_PATH = "./scripts/extraction_audit.csv"
COVERAGE_OUTPUT_PATH = "./scripts/coverage_report.csv"

print("\n" + "=" * 70)
print("PHASE 1 + PHASE 2 : INGESTION COMPLETE")
print("=" * 70 + "\n")


# =====================================================
# HELPERS
# =====================================================

def info_extraction(
    sensitive_columns=None,
    masked_values_count=0,
    free_text_masked_count=0,
    free_text_mask_types=None,
    ocr_pages_count=0,
    ocr_engine="",
):
    """Construit une structure standard d'info d'extraction."""
    return {
        "sensitive_columns": sensitive_columns or [],
        "masked_values_count": masked_values_count,
        "free_text_masked_count": free_text_masked_count,
        "free_text_mask_types": free_text_mask_types or [],
        "ocr_pages_count": ocr_pages_count,
        "ocr_engine": ocr_engine,
    }


def fusionner_infos_extraction(base_info, extra_info):
    """Fusionne deux dictionnaires d'info d'extraction."""
    merged = dict(base_info or {})
    extra_info = extra_info or {}

    for key, value in extra_info.items():
        if key in ("sensitive_columns", "free_text_mask_types"):
            base_values = list(merged.get(key, []))
            for item in value or []:
                if item not in base_values:
                    base_values.append(item)
            merged[key] = base_values
        elif key in ("masked_values_count", "free_text_masked_count", "ocr_pages_count"):
            merged[key] = int(merged.get(key, 0)) + int(value or 0)
        elif key == "ocr_engine":
            if value:
                merged[key] = value
        else:
            merged[key] = value

    return merged


def normaliser_valeur_cellule(value):
    """Convertit une valeur brute en texte propre."""
    if value is None:
        return ""
    return str(value).strip()


def detecter_delimiteur_csv(texte_csv):
    """Detecte le delimiteur CSV via csv.Sniffer + scoring de coherence."""
    sample = texte_csv[:20000]
    candidats = [";", ",", "\t", "|"]

    def score_delimiteur(delimiteur):
        reader = csv.reader(io.StringIO(sample), delimiter=delimiteur)
        tailles = []
        for idx, row in enumerate(reader):
            if idx >= 40:
                break
            if any(normaliser_valeur_cellule(cell) for cell in row):
                tailles.append(len(row))

        if not tailles:
            return -1.0

        multi_colonnes = sum(1 for t in tailles if t > 1)
        moyenne = sum(tailles) / len(tailles)
        stabilite = 1.0 / len(set(tailles))
        return (multi_colonnes * 10.0) + moyenne + stabilite

    sniffer_delim = None
    try:
        sniffer_delim = csv.Sniffer().sniff(sample, delimiters=",;\t|").delimiter
    except Exception:
        pass

    scores = {d: score_delimiteur(d) for d in candidats}
    best_delim = max(scores, key=scores.get)

    if (
        sniffer_delim in scores
        and scores[sniffer_delim] >= 0
        and (scores[sniffer_delim] + 0.5) >= scores[best_delim]
    ):
        return sniffer_delim, "sniffer"

    return best_delim, "scoring"


def extraire_texte_image_ocr(image):
    """OCR image : Tesseract (prioritaire) -> EasyOCR (fallback)."""
    global _EASYOCR_READER

    # Tesseract
    if pytesseract is not None:
        for lang in ["fra+eng", "fra", "eng", None]:
            try:
                texte = (
                    pytesseract.image_to_string(image)
                    if lang is None
                    else pytesseract.image_to_string(image, lang=lang)
                )
                if texte and texte.strip():
                    return texte, f"tesseract_{lang or 'default'}", ""
            except Exception:
                continue

    # EasyOCR fallback
    if easyocr is not None:
        if _EASYOCR_READER is None:
            try:
                _EASYOCR_READER = easyocr.Reader(["fr", "en"], gpu=False)
            except Exception:
                pass
        if _EASYOCR_READER is not None:
            try:
                import numpy as np
                results = _EASYOCR_READER.readtext(np.array(image), detail=0, paragraph=True)
                texte = "\n".join(str(r).strip() for r in results if str(r).strip())
                if texte.strip():
                    return texte, "easyocr", ""
            except Exception:
                pass

    return "", "", "ocr_unavailable"





def lire_pdf_via_pymupdf(filepath):
    """Extraction PDF texte via PyMuPDF (sans OCR)."""
    if fitz is None:
        return "", "", info_extraction()
    try:
        doc = fitz.open(filepath)
        pages = []
        for page in doc:
            t = page.get_text("text") or ""
            if t.strip():
                pages.append(t)
        doc.close()
        texte = "\n".join(pages)
        if texte.strip():
            print("     Extraction via PyMuPDF")
            return texte, "pymupdf", info_extraction()
        return "", "", info_extraction()
    except Exception:
        return "", "", info_extraction()


def masquer_donnees_sensibles_texte_libre(texte):
    """Masque PII/secrets en texte libre (email, tel, tokens, passwords)."""
    texte_masque = texte
    masked_count_total = 0
    masked_types = []

    def appliquer_regex(pattern, mask_type, replace_fn):
        nonlocal texte_masque, masked_count_total
        local_count = 0

        def _repl(match):
            nonlocal local_count
            local_count += 1
            return replace_fn(match)

        texte_masque = pattern.sub(_repl, texte_masque)
        if local_count > 0:
            masked_count_total += local_count
            masked_types.append(f"{mask_type}:{local_count}")

    appliquer_regex(
        PASSWORD_ASSIGN_PATTERN,
        "password_assign",
        lambda m: f"{m.group(1)}{m.group(2)}{SENSITIVE_VALUE_MASK}",
    )
    appliquer_regex(
        SECRET_ASSIGN_PATTERN,
        "secret_assign",
        lambda m: f"{m.group(1)}{m.group(2)}{SENSITIVE_VALUE_MASK}",
    )
    appliquer_regex(EMAIL_PATTERN, "email", lambda _m: "[REDACTED_EMAIL]")
    appliquer_regex(PHONE_PATTERN, "phone", lambda _m: "[REDACTED_PHONE]")
    appliquer_regex(JWT_PATTERN, "jwt", lambda _m: "[REDACTED_JWT]")

    return texte_masque, {
        "free_text_masked_count": masked_count_total,
        "free_text_mask_types": masked_types,
    }


def normaliser_header(header):
    """Normalise un header pour detection de champs sensibles."""
    text = normaliser_valeur_cellule(header).lower()
    return re.sub(r"[^a-z0-9]", "", text)


def header_est_sensible(header):
    """Detecte si un nom de colonne est sensible."""
    norm = normaliser_header(header)
    return any(keyword in norm for keyword in SENSITIVE_COLUMN_KEYWORDS)


def detecter_colonnes_sensibles(headers):
    """Retourne les indices et noms des colonnes sensibles."""
    sensitive_indices = set()
    sensitive_names = []

    for idx, header in enumerate(headers):
        if header_est_sensible(header):
            sensitive_indices.add(idx)
            sensitive_names.append(normaliser_valeur_cellule(header) or f"col_{idx}")

    return sensitive_indices, sensitive_names


def masquer_cellules_sensibles(cells, sensitive_indices):
    """Masque les cellules sensibles d'une ligne."""
    masked = []
    masked_count = 0

    for idx, cell in enumerate(cells):
        cell_text = normaliser_valeur_cellule(cell)
        if idx in sensitive_indices and cell_text:
            masked.append(SENSITIVE_VALUE_MASK)
            masked_count += 1
        else:
            masked.append(cell_text)

    return masked, masked_count


def lire_texte_avec_fallback_encodage(filepath, encodings=None):
    """Lit un fichier texte avec fallback d'encodage."""
    encodings = encodings or TEXT_FILE_ENCODINGS

    for encoding in encodings:
        try:
            with open(filepath, "r", encoding=encoding) as file:
                return file.read(), encoding, None
        except UnicodeDecodeError:
            continue
        except Exception as e:
            return "", "unknown", e

    try:
        with open(filepath, "rb") as file:
            text = file.read().decode("utf-8", errors="ignore")
        return text, "utf-8-ignore", None
    except Exception as e:
        return "", "unknown", e


def extraire_lignes_table_docx(table):
    """Extrait les lignes textuelles d'un tableau DOCX."""
    lignes = []
    for row in table.rows:
        cellules = [normaliser_valeur_cellule(cell.text).replace("\n", " ") for cell in row.cells]
        cellules = [c for c in cellules if c]
        if cellules:
            lignes.append(" | ".join(cellules))
    return lignes


def supprimer_doublons_en_ordre(elements):
    """Supprime les doublons en conservant l'ordre."""
    vus = set()
    uniques = []

    for element in elements:
        if element not in vus:
            vus.add(element)
            uniques.append(element)

    return uniques


def get_source_relative_path(filepath):
    """Retourne le chemin relatif normalise depuis DOCUMENTS_PATH."""
    return os.path.relpath(filepath, DOCUMENTS_PATH).replace("\\", "/")


def get_categorie(filepath):
    """Retourne la categorie top-level (AFD/CdC/Exigences/Guide...)."""
    source_relative_path = get_source_relative_path(filepath)
    path_lower = source_relative_path.lower()

    if "afd" in path_lower:
        return "AFD"
    if "cdc" in path_lower or "cahier des charges" in path_lower:
        return "CdC"
    if "exigences" in path_lower:
        return "Exigences"
    if "guide" in path_lower:
        return "Guide"

    parts = source_relative_path.split("/")
    if len(parts) > 1:
        return parts[0]

    return "<unknown>"


def collecter_tous_les_fichiers():
    """Collecte tous les fichiers du corpus (hors lock files Office)."""
    fichiers = []

    for root, _dirs, files in os.walk(DOCUMENTS_PATH):
        for file_name in files:
            if file_name.startswith("~$"):
                continue
            filepath = os.path.join(root, file_name)
            extension = os.path.splitext(file_name)[1].lower()
            fichiers.append(
                {
                    "filepath": filepath,
                    "source_relative_path": get_source_relative_path(filepath),
                    "category": get_categorie(filepath),
                    "extension": extension,
                }
            )

    return fichiers


def calculer_taux(part, total):
    """Calcule un pourcentage formate."""
    if total == 0:
        return "0.0%"
    return f"{(part / total) * 100:.1f}%"


def apercu_texte(texte, max_chars=500):
    """Retourne un apercu compact d'un texte pour les logs."""
    if not texte:
        return ""

    compact = re.sub(r"\s+", " ", texte).strip()
    if len(compact) <= max_chars:
        return compact

    return compact[:max_chars] + " ..."


def logger_resultat_etape(nom_etape, texte, max_chars=None):
    """Affiche un extrait lisible du resultat d'etape."""
    if not PRINT_STEP_RESULTS:
        return

    max_chars = max_chars or STEP_PREVIEW_CHARS
    print(f"  [RESULTAT] {nom_etape}")
    print(f"     Longueur: {len(texte)} caracteres")
    print(f"     Apercu: {apercu_texte(texte, max_chars=max_chars)}")


def ajouter_entree_audit(
    audit_entries,
    filepath,
    type_doc,
    extraction_method,
    extraction_info,
    raw_length,
    cleaned_length,
    status,
    indexed,
):
    """Ajoute une entree d'audit d'extraction."""
    source_relative_path = get_source_relative_path(filepath)
    sensitive_columns = extraction_info.get("sensitive_columns", [])
    free_text_mask_types = extraction_info.get("free_text_mask_types", [])

    audit_entries.append(
        {
            "timestamp": datetime.now().isoformat(),
            "source": os.path.basename(filepath),
            "source_relative_path": source_relative_path,
            "category": get_categorie(filepath),
            "type": type_doc,
            "extension": os.path.splitext(filepath)[1].lower(),
            "extraction_method": extraction_method,
            "raw_length": raw_length,
            "cleaned_length": cleaned_length,
            "status": status,
            "indexed": indexed,
            "sensitive_columns": "|".join(sensitive_columns),
            "masked_values_count": extraction_info.get("masked_values_count", 0),
            "free_text_masked_count": extraction_info.get("free_text_masked_count", 0),
            "free_text_mask_types": "|".join(free_text_mask_types),
            "ocr_pages_count": extraction_info.get("ocr_pages_count", 0),
            "ocr_engine": extraction_info.get("ocr_engine", ""),
        }
    )


def ecrire_rapport_audit(audit_entries):
    """Ecrit le rapport d'audit et affiche les KPI d'extraction."""
    if not audit_entries:
        print("\nAucun audit a ecrire")
        return

    os.makedirs(os.path.dirname(AUDIT_OUTPUT_PATH), exist_ok=True)

    fieldnames = [
        "timestamp",
        "source",
        "source_relative_path",
        "category",
        "type",
        "extension",
        "extraction_method",
        "raw_length",
        "cleaned_length",
        "status",
        "indexed",
        "sensitive_columns",
        "masked_values_count",
        "free_text_masked_count",
        "free_text_mask_types",
        "ocr_pages_count",
        "ocr_engine",
    ]

    with open(AUDIT_OUTPUT_PATH, "w", encoding="utf-8", newline="") as csv_file:
        writer = csv.DictWriter(csv_file, fieldnames=fieldnames, delimiter=";")
        writer.writeheader()
        writer.writerows(audit_entries)

    total = len(audit_entries)
    indexed_count = sum(1 for e in audit_entries if e["indexed"])
    low_text_count = sum(1 for e in audit_entries if e["status"] == "low_text")
    empty_count = sum(1 for e in audit_entries if e["status"].startswith("empty"))
    error_count = sum(1 for e in audit_entries if e["status"] == "error")
    masked_docs = sum(1 for e in audit_entries if e["masked_values_count"] > 0)
    free_text_mask_total = sum(int(e["free_text_masked_count"]) for e in audit_entries)
    ocr_pages_total = sum(int(e["ocr_pages_count"]) for e in audit_entries)

    print("\n" + "=" * 70)
    print("AUDIT EXTRACTION")
    print("=" * 70)
    print(f"Documents audites: {total}")
    print(f"Documents indexes: {indexed_count} ({calculer_taux(indexed_count, total)})")
    print(f"Documents low text: {low_text_count} ({calculer_taux(low_text_count, total)})")
    print(f"Documents vides: {empty_count} ({calculer_taux(empty_count, total)})")
    print(f"Erreurs extraction/pipeline: {error_count} ({calculer_taux(error_count, total)})")
    print(f"Docs avec masquage sensible: {masked_docs} ({calculer_taux(masked_docs, total)})")
    print(f"Total masquages texte libre: {free_text_mask_total}")
    print(f"Pages OCR traitees: {ocr_pages_total}")
    print(f"Rapport detaille: {AUDIT_OUTPUT_PATH}")


def generer_rapport_couverture(audit_entries):
    """Genere un rapport de couverture par categorie et extension."""
    tous_les_fichiers = collecter_tous_les_fichiers()
    total_files = len(tous_les_fichiers)

    indexed_paths = {
        entry["source_relative_path"]
        for entry in audit_entries
        if entry.get("indexed")
    }

    total_by_category = {}
    indexed_by_category = {}
    total_by_extension = {}
    indexed_by_extension = {}

    for file_info in tous_les_fichiers:
        category = file_info["category"]
        extension = file_info["extension"] or "<none>"
        rel_path = file_info["source_relative_path"]

        total_by_category[category] = total_by_category.get(category, 0) + 1
        total_by_extension[extension] = total_by_extension.get(extension, 0) + 1

        if rel_path in indexed_paths:
            indexed_by_category[category] = indexed_by_category.get(category, 0) + 1
            indexed_by_extension[extension] = indexed_by_extension.get(extension, 0) + 1

    indexed_total = len(indexed_paths)
    missing_total = total_files - indexed_total

    print("\n" + "=" * 70)
    print("RAPPORT DE COUVERTURE")
    print("=" * 70)
    print("Formule: coverage = indexed_files / total_files")
    print(f"Global: {indexed_total}/{total_files} ({calculer_taux(indexed_total, total_files)})")
    print(f"Missing: {missing_total}")

    print("\nCouverture par categorie:")
    for category in sorted(total_by_category.keys()):
        total = total_by_category[category]
        indexed = indexed_by_category.get(category, 0)
        missing = total - indexed
        print(
            f"- {category}: {indexed}/{total} ({calculer_taux(indexed, total)}) | missing: {missing}"
        )

    print("\nCouverture par extension:")
    for extension in sorted(total_by_extension.keys()):
        total = total_by_extension[extension]
        indexed = indexed_by_extension.get(extension, 0)
        missing = total - indexed
        print(
            f"- {extension}: {indexed}/{total} ({calculer_taux(indexed, total)}) | missing: {missing}"
        )

    os.makedirs(os.path.dirname(COVERAGE_OUTPUT_PATH), exist_ok=True)
    fieldnames = [
        "timestamp",
        "scope",
        "name",
        "total_files",
        "indexed_files",
        "missing_files",
        "coverage_pct",
    ]

    rows = [
        {
            "timestamp": datetime.now().isoformat(),
            "scope": "overall",
            "name": "all",
            "total_files": total_files,
            "indexed_files": indexed_total,
            "missing_files": missing_total,
            "coverage_pct": calculer_taux(indexed_total, total_files),
        }
    ]

    for category in sorted(total_by_category.keys()):
        total = total_by_category[category]
        indexed = indexed_by_category.get(category, 0)
        rows.append(
            {
                "timestamp": datetime.now().isoformat(),
                "scope": "category",
                "name": category,
                "total_files": total,
                "indexed_files": indexed,
                "missing_files": total - indexed,
                "coverage_pct": calculer_taux(indexed, total),
            }
        )

    for extension in sorted(total_by_extension.keys()):
        total = total_by_extension[extension]
        indexed = indexed_by_extension.get(extension, 0)
        rows.append(
            {
                "timestamp": datetime.now().isoformat(),
                "scope": "extension",
                "name": extension,
                "total_files": total,
                "indexed_files": indexed,
                "missing_files": total - indexed,
                "coverage_pct": calculer_taux(indexed, total),
            }
        )

    with open(COVERAGE_OUTPUT_PATH, "w", encoding="utf-8", newline="") as csv_file:
        writer = csv.DictWriter(csv_file, fieldnames=fieldnames, delimiter=";")
        writer.writeheader()
        writer.writerows(rows)

    print(f"\nRapport couverture: {COVERAGE_OUTPUT_PATH}")
    print("Conseil: lance ce script chaque semaine et compare coverage_report.csv")


# =====================================================
# PHASE 1 : INGESTION
# =====================================================

print("PHASE 1 : INGESTION\n")


def lire_pdf(filepath):
    """Extrait le texte d'un PDF : PyMuPDF -> pdfplumber -> PyPDF2 (sans OCR)."""
    print(f"  Lecture PDF: {os.path.basename(filepath)}")

    # 1) PyMuPDF
    texte, method, info = lire_pdf_via_pymupdf(filepath)
    if texte.strip():
        return texte, method, info

    # 2) pdfplumber
    if pdfplumber is not None:
        try:
            with pdfplumber.open(filepath) as pdf:
                pages_text = []
                tables_pages_count = 0

                for page in pdf.pages:
                    extracted = page.extract_text() or ""
                    parts = []
                    if extracted.strip():
                        parts.append(extracted)

                    if PDF_EXTRACT_TABLES:
                        try:
                            table_lines = []
                            for table in (page.extract_tables() or []):
                                for row in (table or []):
                                    row_clean = [normaliser_valeur_cellule(c) for c in (row or [])]
                                    if any(row_clean):
                                        table_lines.append(" | ".join(row_clean))
                            if table_lines:
                                parts.append("\n".join(table_lines))
                                tables_pages_count += 1
                        except Exception:
                            pass

                    if parts:
                        pages_text.append("\n".join(parts))

                texte = "\n".join(pages_text)
                if texte.strip():
                    method = "pdfplumber"
                    if tables_pages_count > 0:
                        method += "_plus_tables"
                    print("     Extraction via pdfplumber")
                    return texte, method, info_extraction()

                print("     pdfplumber vide, fallback PyPDF2...")
        except Exception as e:
            print(f"     Erreur pdfplumber: {e}")

    # 3) PyPDF2
    texte = ""
    try:
        with open(filepath, "rb") as file:
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    texte += extracted + "\n"
        if texte.strip():
            print("     Extraction via PyPDF2")
        return texte, "pypdf2_fallback", info_extraction()
    except Exception as e:
        print(f"     Erreur PDF: {e}")
        return "", "pdf_error", info_extraction()


def lire_docx(filepath, category=None):
    """Extrait le texte d'un DOCX (paragraphes + tableaux + en-tetes/pieds).
    Ouvre en mode binaire pour eviter les erreurs de chemin Windows (apostrophes).
    Images embarquees : description textuelle pour les images > 50KB.
    """
    print(f"  Lecture DOCX: {os.path.basename(filepath)}")

    try:
        # open() binaire evite les bugs de chemin Windows avec python-docx
        with open(filepath, 'rb') as f:
            doc = Document(f)

        paragraphes = [normaliser_valeur_cellule(p.text) for p in doc.paragraphs if normaliser_valeur_cellule(p.text)]

        tableaux = []
        for table in doc.tables:
            tableaux.extend(extraire_lignes_table_docx(table))

        entetes_pieds = []
        for section in doc.sections:
            entetes_pieds.extend(
                [normaliser_valeur_cellule(p.text) for p in section.header.paragraphs if normaliser_valeur_cellule(p.text)]
            )
            entetes_pieds.extend(
                [normaliser_valeur_cellule(p.text) for p in section.footer.paragraphs if normaliser_valeur_cellule(p.text)]
            )
            for table in section.header.tables:
                entetes_pieds.extend(extraire_lignes_table_docx(table))
            for table in section.footer.tables:
                entetes_pieds.extend(extraire_lignes_table_docx(table))

        entetes_pieds = supprimer_doublons_en_ordre(entetes_pieds)

        blocs = []
        if paragraphes:
            blocs.append("\n".join(paragraphes))
        if tableaux:
            blocs.append("\n".join(tableaux))
        if entetes_pieds:
            blocs.append("\n".join(entetes_pieds))

        texte = "\n".join(blocs)
        print(f"     Paragraphes: {len(paragraphes)}")
        print(f"     Lignes tableaux: {len(tableaux)}")
        print(f"     En-tetes/pieds: {len(entetes_pieds)}")

        extraction_info = info_extraction()

        # Images embarquees : description textuelle pour les images > 50KB
        if Image is not None:
            try:
                images_descriptions = []
                images_count = 0

                with zipfile.ZipFile(filepath) as zf:
                    for name in zf.namelist():
                        lower = name.lower()
                        if not lower.startswith("word/media/"):
                            continue
                        if lower.endswith((".emf", ".wmf")):
                            continue
                        if not lower.endswith((".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff")):
                            continue

                        data = zf.read(name)
                        # Ignorer les petites images (icones, puces) < 50KB
                        if len(data) < 50_000:
                            continue

                        nom_image = Path(name).stem.replace("_", " ").replace("-", " ")
                        desc = (
                            f"[Capture ecran Smart Factory: {nom_image}] "
                            f"Interface utilisateur de l'application Smart Factory MOMsoft."
                        )
                        images_descriptions.append(desc)
                        images_count += 1

                if images_descriptions:
                    texte = "\n\n".join([texte, "\n\n".join(images_descriptions)]) if texte.strip() else "\n\n".join(images_descriptions)
                    print(f"     Images DOCX describees: {images_count} image(s) > 50KB")
            except Exception:
                pass

        method = "docx_text"
        if extraction_info.get("ocr_pages_count", 0) > 0:
            method += "_plus_ocr"

        return texte, method, extraction_info
    except Exception as e:
        print(f"     Erreur DOCX: {e}")
        return "", "docx_error", info_extraction()


def lire_txt(filepath):
    """Lit un fichier TXT avec fallback d'encodage."""
    print(f"  Lecture TXT: {os.path.basename(filepath)}")
    texte, encoding, error = lire_texte_avec_fallback_encodage(filepath)

    if error is not None:
        print(f"     Erreur TXT: {error}")
        return "", "txt_error", info_extraction()

    print(f"     Encodage: {encoding}")
    return texte, f"txt_{encoding}", info_extraction()


def lire_csv(filepath):
    """Lit un fichier CSV en masquant les colonnes sensibles."""
    print(f"  Lecture CSV: {os.path.basename(filepath)}")

    texte_csv, encoding, error = lire_texte_avec_fallback_encodage(filepath)
    if error is not None:
        print(f"     Erreur CSV: {error}")
        return "", "csv_error", info_extraction()

    delimiter, delimiter_detection_method = detecter_delimiteur_csv(texte_csv)
    delimiter_labels = {";": "semicolon", ",": "comma", "\t": "tab", "|": "pipe"}
    delimiter_label = delimiter_labels.get(delimiter, "other")
    print(f"     Delimiteur detecte: {repr(delimiter)} ({delimiter_detection_method})")

    reader = csv.reader(io.StringIO(texte_csv), delimiter=delimiter)
    rows = list(reader)
    if not rows:
        return "", f"csv_{encoding}_empty", info_extraction()

    headers = [normaliser_valeur_cellule(h) for h in rows[0]]
    sensitive_indices, sensitive_columns = detecter_colonnes_sensibles(headers)
    masked_values_count = 0

    if sensitive_columns:
        print(f"     Colonnes sensibles detectees: {', '.join(sensitive_columns)}")
        print("     Politique: valeurs sensibles masquees")

    lignes = []
    if headers:
        lignes.append(" | ".join(headers))

    for row in rows[1:]:
        masked_row, masked_count = masquer_cellules_sensibles(row, sensitive_indices)
        masked_values_count += masked_count
        if any(masked_row):
            lignes.append(" | ".join(masked_row))

    texte = "\n".join(lignes)
    extraction_method = f"csv_{encoding}_{delimiter_detection_method}_{delimiter_label}"
    if sensitive_columns:
        extraction_method += "_masked"

    return texte, extraction_method, info_extraction(sensitive_columns, masked_values_count)


def lire_xlsx(filepath):
    """Lit un fichier XLSX en masquant les colonnes sensibles."""
    print(f"  Lecture XLSX: {os.path.basename(filepath)}")

    if openpyxl is None:
        print("     openpyxl non installe, fallback nom du fichier")
        fallback = f"Fichier XLSX: {Path(filepath).stem.replace('_', ' ').replace('-', ' ')}"
        return fallback, "xlsx_filename_fallback_missing_openpyxl", info_extraction()

    try:
        workbook = openpyxl.load_workbook(filepath, read_only=True, data_only=True)
        lignes = []
        global_sensitive_columns = set()
        masked_values_count = 0

        for sheet in workbook.worksheets:
            rows_iter = sheet.iter_rows(values_only=True)
            first_row = next(rows_iter, None)
            if first_row is None:
                continue

            headers = [normaliser_valeur_cellule(v) for v in first_row]
            sensitive_indices, sensitive_columns = detecter_colonnes_sensibles(headers)
            global_sensitive_columns.update(sensitive_columns)

            lignes.append(f"[Sheet: {sheet.title}]")
            if any(headers):
                lignes.append(" | ".join(headers))

            for row in rows_iter:
                masked_row, masked_count = masquer_cellules_sensibles(row, sensitive_indices)
                masked_values_count += masked_count
                if any(masked_row):
                    lignes.append(" | ".join(masked_row))

        workbook.close()

        sensitive_columns_list = sorted(global_sensitive_columns)
        if sensitive_columns_list:
            print(f"     Colonnes sensibles detectees: {', '.join(sensitive_columns_list)}")
            print("     Politique: valeurs sensibles masquees")

        texte = "\n".join(lignes)
        method = "xlsx_openpyxl"
        if sensitive_columns_list:
            method += "_masked"

        return texte, method, info_extraction(sensitive_columns_list, masked_values_count)
    except Exception as e:
        print(f"     Erreur XLSX: {e}")
        return "", "xlsx_error", info_extraction()


def lire_html(filepath):
    """Lit un fichier HTML et extrait le texte visible."""
    print(f"  Lecture HTML: {os.path.basename(filepath)}")

    html_content, encoding, error = lire_texte_avec_fallback_encodage(filepath)
    if error is not None:
        print(f"     Erreur HTML: {error}")
        return "", "html_error", info_extraction()

    if BeautifulSoup is not None:
        try:
            soup = BeautifulSoup(html_content, "html.parser")
            for tag in soup(["script", "style", "noscript"]):
                tag.decompose()

            texte = soup.get_text("\n")
            print("     Extraction via BeautifulSoup")
            return texte, f"html_bs4_{encoding}", info_extraction()
        except Exception as e:
            print(f"     Erreur BeautifulSoup: {e}")

    # Fallback regex simple
    texte = re.sub(r"<[^>]+>", " ", html_content)
    texte = unescape(texte)
    print("     Extraction HTML regex fallback")
    return texte, f"html_regex_{encoding}", info_extraction()


def extraire_texte_xml(root):
    """Extrait texte + attributs d'un arbre XML."""
    morceaux = []

    for node in root.iter():
        node_text = normaliser_valeur_cellule(node.text)
        if node_text:
            morceaux.append(node_text)

        for attr_name, attr_value in node.attrib.items():
            value_text = normaliser_valeur_cellule(attr_value)
            if value_text:
                morceaux.append(f"{attr_name}: {value_text}")

    return supprimer_doublons_en_ordre(morceaux)


def lire_jrxml(filepath):
    """Lit un fichier JRXML/XML."""
    print(f"  Lecture XML/JRXML: {os.path.basename(filepath)}")

    try:
        root = ET.parse(filepath).getroot()
        morceaux = extraire_texte_xml(root)
        texte = "\n".join(morceaux)
        return texte, "jrxml_xml_parser", info_extraction()
    except ET.ParseError:
        print("     Parse XML impossible, fallback texte brut")
    except Exception as e:
        print(f"     Erreur XML: {e}")

    text_content, encoding, error = lire_texte_avec_fallback_encodage(filepath)
    if error is not None:
        print(f"     Erreur fallback XML: {error}")
        return "", "jrxml_error", info_extraction()

    texte = re.sub(r"<[^>]+>", " ", text_content)
    texte = unescape(texte)
    return texte, f"jrxml_text_fallback_{encoding}", info_extraction()


def lire_image(filepath):
    """Cree une description textuelle de l'image depuis son nom de fichier.
    Les screenshots Smart Factory produisent du bruit en OCR, on prefere une description propre."""
    print(f"  Lecture image: {os.path.basename(filepath)}")

    nom = Path(filepath).stem.replace("_", " ").replace("-", " ")
    description = (
        f"Capture ecran Smart Factory: {nom}. "
        f"Interface utilisateur de l'application Smart Factory MOMsoft. "
        f"Source: {get_source_relative_path(filepath)}"
    )
    return description, "image_description_from_filename", info_extraction()


def lire_document(filepath):
    """Lit un document selon son extension.

    Retourne: texte, extraction_method, extraction_info
    """
    extension = os.path.splitext(filepath)[1].lower()
    category = get_categorie(filepath)

    if extension == ".pdf":
        return lire_pdf(filepath)
    if extension == ".docx":
        return lire_docx(filepath, category=category)
    if extension == ".txt":
        return lire_txt(filepath)
    if extension == ".csv":
        return lire_csv(filepath)
    if extension == ".xlsx":
        return lire_xlsx(filepath)
    if extension in (".html", ".htm"):
        return lire_html(filepath)
    if extension in (".jrxml", ".xml"):
        return lire_jrxml(filepath)
    if extension in (".png", ".jpg", ".jpeg"):
        return lire_image(filepath)

    print(f"  Format non supporte: {extension}")
    return "", f"unsupported_{extension}", info_extraction()


# ===== 2. NETTOYAGE =====

def nettoyer_texte(texte):
    """Nettoie le texte extrait."""
    print("  Nettoyage du texte...")

    # Supprimer lignes vides multiples
    lignes = texte.split("\n")
    lignes = [ligne.strip() for ligne in lignes if ligne.strip()]
    texte_propre = "\n".join(lignes)

    # Supprimer caracteres speciaux problematiques
    texte_propre = texte_propre.replace("\x00", "")
    texte_propre = texte_propre.replace("\uf0b7", "•")

    # Supprimer espaces multiples
    texte_propre = re.sub(r"[ \t]+", " ", texte_propre)
    texte_propre = re.sub(r"\n{3,}", "\n\n", texte_propre)

    print(f"     Avant: {len(texte)} caracteres")
    print(f"     Apres: {len(texte_propre)} caracteres")
    return texte_propre


# ===== 3. CHUNKING (DECOUPAGE) =====

def decouper_en_chunks(texte, metadata):
    """Decoupe le texte en chunks par paragraphe puis par phrase (sans LangChain)."""
    print("  Decoupage en chunks...")

    # Separateurs par priorite : double saut de ligne, simple saut, point-espace, espace
    paragraphes = texte.split("\n\n")
    chunks = []
    current = ""

    for para in paragraphes:
        para = para.strip()
        if not para:
            continue
        # Si le paragraphe tient dans un chunk avec le buffer courant
        candidate = (current + "\n\n" + para).strip() if current else para
        if len(candidate) <= CHUNK_SIZE:
            current = candidate
        else:
            # Flush le buffer courant
            if current:
                chunks.append(current)
            # Si le paragraphe seul depasse CHUNK_SIZE, decouper par phrase
            if len(para) > CHUNK_SIZE:
                phrases = re.split(r'(?<=[.!?])\s+', para)
                sub = ""
                for phrase in phrases:
                    candidate_sub = (sub + " " + phrase).strip() if sub else phrase
                    if len(candidate_sub) <= CHUNK_SIZE:
                        sub = candidate_sub
                    else:
                        if sub:
                            chunks.append(sub)
                        # phrase trop longue: couper par mots
                        if len(phrase) > CHUNK_SIZE:
                            words = phrase.split()
                            sub = ""
                            for w in words:
                                cand = (sub + " " + w).strip() if sub else w
                                if len(cand) <= CHUNK_SIZE:
                                    sub = cand
                                else:
                                    if sub:
                                        chunks.append(sub)
                                    sub = w
                        else:
                            sub = phrase
                if sub:
                    current = sub
                else:
                    current = ""
            else:
                current = para

    if current:
        chunks.append(current)

    # Appliquer l'overlap : reprendre les derniers CHUNK_OVERLAP caracteres du chunk precedent
    if CHUNK_OVERLAP > 0 and len(chunks) > 1:
        overlapped = [chunks[0]]
        for i in range(1, len(chunks)):
            prev = chunks[i - 1]
            overlap_text = prev[-CHUNK_OVERLAP:] if len(prev) > CHUNK_OVERLAP else prev
            # Couper au premier espace pour ne pas couper un mot
            space_idx = overlap_text.find(" ")
            if space_idx > 0:
                overlap_text = overlap_text[space_idx + 1:]
            overlapped.append((overlap_text + " " + chunks[i]).strip())
        chunks = overlapped

    if not chunks:
        print("     Aucun chunk produit")
        return []

    chunks_avec_metadata = []
    for i, chunk in enumerate(chunks):
        chunk_data = {
            "texte": chunk,
            "metadata": {
                **metadata,
                "chunk_id": i,
                "total_chunks": len(chunks),
            },
        }
        chunks_avec_metadata.append(chunk_data)

    print(f"     {len(chunks)} chunks crees")
    print(f"     Taille moyenne: {sum(len(c) for c in chunks) // len(chunks)} caracteres")
    return chunks_avec_metadata


# =====================================================
# PHASE 2 : EMBEDDINGS + STOCKAGE
# =====================================================

print("\nPHASE 2 : EMBEDDINGS & STOCKAGE\n")


def creer_modele_embedding():
    """Cree le modele d'embedding."""
    print(f"Chargement du modele d'embedding: {EMBEDDING_MODEL_NAME}")
    try:
        from sentence_transformers import SentenceTransformer
    except Exception as e:
        raise RuntimeError(
            "Impossible de charger sentence-transformers. "
            "Mets a jour les dependances avec: pip install -r requirements.txt"
        ) from e

    model = SentenceTransformer(EMBEDDING_MODEL_NAME)
    print(f"   Modele charge (dim={model.get_sentence_embedding_dimension()})")
    return model


def initialiser_vector_db():
    """Initialise ChromaDB."""
    print("Initialisation de la base vectorielle ChromaDB...")

    if chromadb is None:
        raise RuntimeError(
            "chromadb n'est pas installe. Lance: pip install -r requirements.txt"
        )

    client = chromadb.PersistentClient(path=VECTOR_DB_PATH)

    try:
        collection = client.get_collection(name="documents")
        print("   Collection existante recuperee")
        print(f"   Contient deja {collection.count()} chunks")
    except Exception:
        collection = client.create_collection(
            name="documents",
            metadata={"hnsw:space": "cosine"},
        )
        print("   Nouvelle collection creee")

    return client, collection


def construire_chunk_id(metadata):
    """Construit un ID stable: {category}_{sha1(source)[:8]}_{chunk_id:04d}."""
    source = metadata.get("source", "unknown")
    category = metadata.get("category", "doc")
    digest = hashlib.sha1(source.encode("utf-8")).hexdigest()[:8]
    chunk_idx = metadata.get("chunk_id", 0)
    return f"{category}_{digest}_{chunk_idx:04d}"


def ajouter_chunks_db(chunks, model, collection):
    """Ajoute les chunks dans ChromaDB avec deduplication par ID."""
    print(f"  Ajout de {len(chunks)} chunks dans ChromaDB...")

    textes = [chunk["texte"] for chunk in chunks]
    metadatas = [chunk["metadata"] for chunk in chunks]
    ids = [construire_chunk_id(chunk["metadata"]) for chunk in chunks]

    existing_ids = set()
    try:
        existing = collection.get(ids=ids)
        existing_ids = set(existing["ids"])
    except Exception:
        pass

    nouveaux_indices = [i for i, id_ in enumerate(ids) if id_ not in existing_ids]
    if not nouveaux_indices:
        print("     Tous les chunks existent deja, ignore")
        return 0

    t_new = [textes[i] for i in nouveaux_indices]
    m_new = [metadatas[i] for i in nouveaux_indices]
    id_new = [ids[i] for i in nouveaux_indices]

    print(f"     {len(nouveaux_indices)} nouveaux chunks -> ChromaDB")
    print("     Generation des embeddings...")
    # Prefixe 'passage:' requis uniquement par multilingual-e5 (asymetric retrieval)
    prefix = "passage: " if "e5" in EMBEDDING_MODEL_NAME.lower() else ""
    t_prefixed = [f"{prefix}{t}" for t in t_new]
    vecteurs = model.encode(t_prefixed, show_progress_bar=True).tolist()
    collection.add(embeddings=vecteurs, documents=t_new, metadatas=m_new, ids=id_new)
    print(f"     Chunks ajoutes avec succes -> {collection.count()} total")
    return len(nouveaux_indices)


def ingerer_document(filepath, type_doc, model, collection):
    """Pipeline complet d'ingestion d'un document.

    Retourne:
    indexed, extraction_method, extraction_info, raw_len, cleaned_len, status, chunks
    """
    print(f"\n{'=' * 70}")
    print(f"Document: {os.path.basename(filepath)}")
    print(f"Type: {type_doc}")
    print(f"{'=' * 70}\n")

    # Lecture
    texte_brut, extraction_method, extraction_info = lire_document(filepath)
    raw_len = len(texte_brut)
    logger_resultat_etape("Extraction", texte_brut)

    if not texte_brut or not texte_brut.strip():
        print("  Document vide apres extraction")
        return False, extraction_method, extraction_info, raw_len, 0, "empty", []

    # Nettoyage
    texte_propre = nettoyer_texte(texte_brut)

    # Masquage PII / secrets en texte libre (pas seulement colonnes tabulaires)
    texte_propre, free_text_info = masquer_donnees_sensibles_texte_libre(texte_propre)
    extraction_info = fusionner_infos_extraction(extraction_info, free_text_info)
    if extraction_info.get("free_text_masked_count", 0) > 0:
        print(
            "  Masquage texte libre applique: "
            f"{extraction_info.get('free_text_masked_count', 0)} occurrence(s)"
        )

    cleaned_len = len(texte_propre)
    logger_resultat_etape("Nettoyage", texte_propre)

    if not texte_propre or not texte_propre.strip():
        print("  Document vide apres nettoyage")
        return False, extraction_method, extraction_info, raw_len, cleaned_len, "empty_after_cleaning", []

    # low_text est audite mais le document reste indexe
    status = "ok"
    if cleaned_len < LOW_TEXT_THRESHOLD:
        print("  Warning: document low text (< threshold), indexation maintenue")
        status = "low_text"

    metadata = {
        "source": os.path.basename(filepath),
        "source_relative_path": get_source_relative_path(filepath),
        "category": get_categorie(filepath),
        "type": type_doc,
        "date_ajout": datetime.now().isoformat(),
        "extraction_method": extraction_method,
        "sensitive_masking_applied": bool(extraction_info.get("sensitive_columns"))
        or extraction_info.get("free_text_masked_count", 0) > 0,
        "sensitive_columns": "|".join(extraction_info.get("sensitive_columns", [])),
        "masked_values_count": extraction_info.get("masked_values_count", 0),
        "free_text_masked_count": extraction_info.get("free_text_masked_count", 0),
        "free_text_mask_types": "|".join(extraction_info.get("free_text_mask_types", [])),
        "ocr_pages_count": extraction_info.get("ocr_pages_count", 0),
        "ocr_engine": extraction_info.get("ocr_engine", ""),
        "embedding_model": EMBEDDING_MODEL_NAME,
    }

    chunks = decouper_en_chunks(texte_propre, metadata)
    if not chunks:
        return False, extraction_method, extraction_info, raw_len, cleaned_len, "no_chunk", []

    if PRINT_STEP_RESULTS:
        print("  [RESULTAT] Chunking")
        print(f"     Total chunks: {len(chunks)}")
        for i, chunk in enumerate(chunks[:STEP_CHUNKS_PREVIEW]):
            chunk_text = chunk.get("texte", "")
            print(
                f"     Chunk {i + 1}/{len(chunks)} (len={len(chunk_text)}): "
                f"{apercu_texte(chunk_text, max_chars=STEP_PREVIEW_CHARS)}"
            )

    # Mode phase 1 only: lecture + nettoyage + chunking sans embeddings/vector DB.
    if model is None:
        return True, extraction_method, extraction_info, raw_len, cleaned_len, "phase1_only", chunks

    added_count = ajouter_chunks_db(chunks, model, collection)
    if added_count == 0:
        # Le document est deja indexe, on considere "indexed" = True.
        return True, extraction_method, extraction_info, raw_len, cleaned_len, "already_indexed", chunks

    return True, extraction_method, extraction_info, raw_len, cleaned_len, status, chunks


def exporter_chunks_vers_csv(all_chunks):
    """Exporte tous les chunks generes dans un fichier CSV avec format spécifique."""
    if not all_chunks:
        return
        
    # 1. Ignorer la catégorie <unknown>
    filtered_chunks = [c for c in all_chunks if c.get("metadata", {}).get("category", "") != "<unknown>"]
    if not filtered_chunks:
        print("\n[!] Aucun chunk à exporter après le filtre (les catégories <unknown> ont été ignorées).")
        return
        
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    csv_path = f"chunks_{timestamp}.csv"
    
    # 2. Définir des colonnes claires
    fieldnames = [
        "id", "source", "source_relative_path", "category", "type", 
        "extension", "extraction_method", "chunk_id", "total_chunks", 
        "date_ajout", "text_len", "text"
    ]
    
    try:
        with open(csv_path, "w", encoding="utf-8", newline="") as csv_file:
            writer = csv.DictWriter(csv_file, fieldnames=fieldnames, delimiter=",")
            writer.writeheader()
            
            for chunk in filtered_chunks:
                meta = chunk.get("metadata", {})
                source = meta.get("source", "doc")
                chunk_id = meta.get("chunk_id", 0)
                
                # 3. Construire un ID clair
                clair_id = f"{source}_{chunk_id}"
                
                # Deduire extension de la source
                extension = os.path.splitext(source)[1].lower() if source else ""
                
                row = {
                    "id": clair_id,
                    "source": source,
                    "source_relative_path": meta.get("source_relative_path", ""),
                    "category": meta.get("category", ""),
                    "type": meta.get("type", ""),
                    "extension": extension,
                    "extraction_method": meta.get("extraction_method", ""),
                    "chunk_id": chunk_id,
                    "total_chunks": meta.get("total_chunks", ""),
                    "date_ajout": meta.get("date_ajout", ""),
                    "text_len": len(chunk.get("texte", "")),
                    "text": chunk.get("texte", "")
                }
                writer.writerow(row)
        print(f"\n[+] Export de {len(filtered_chunks)} chunks réussi vers : {csv_path} (Ignorés: {len(all_chunks) - len(filtered_chunks)} <unknown>)")
    except Exception as e:
        print(f"\n[-] Erreur lors de l'export des chunks: {e}")


def indexer_document(filepath, category, collection=None, model=None):
    """Indexe un seul document dans ChromaDB (appel runtime depuis l'API).

    Args:
        filepath:   chemin absolu du fichier
        category:   CdC | Exigences | Guide | AFD
        collection: collection ChromaDB (si None, initialise)
        model:      modele d'embedding (si None, charge)

    Returns:
        int: nombre de chunks ajoutes
    """
    if model is None:
        model = creer_modele_embedding()
    if collection is None:
        _, collection = initialiser_vector_db()

    type_doc = category
    indexed, extraction_method, extraction_info, raw_len, cleaned_len, status, chunks = \
        ingerer_document(filepath, type_doc, model, collection)

    if indexed and chunks:
        return len(chunks)
    return 0


def ingerer_tous_documents():
    """Ingere tous les documents du dossier."""
    print(f"Scan du dossier: {DOCUMENTS_PATH}\n")

    if not os.path.exists(DOCUMENTS_PATH):
        print(f"Dossier {DOCUMENTS_PATH} introuvable")
        return

    tous_les_fichiers = collecter_tous_les_fichiers()
    fichiers_a_traiter = []

    for file_info in tous_les_fichiers:
        if file_info["extension"] in SUPPORTED_EXTENSIONS:
            filepath = file_info["filepath"]
            type_doc = os.path.basename(os.path.dirname(filepath))
            fichiers_a_traiter.append((filepath, type_doc))

    if not fichiers_a_traiter:
        print("Aucun document a traiter")
        return

    print(f"Total fichiers corpus: {len(tous_les_fichiers)}")
    print(f"Fichiers supportes pour ingestion: {len(fichiers_a_traiter)}")
    print(f"Extensions supportees: {', '.join(SUPPORTED_EXTENSIONS)}")

    model = None
    collection = None

    if RUN_PHASE2:
        try:
            model = creer_modele_embedding()
            _client, collection = initialiser_vector_db()
        except Exception as e:
            print("\n" + "!" * 70)
            print("PHASE 2 indisponible -> bascule automatique en mode PHASE 1 ONLY")
            print(f"Raison: {e}")
            print("Aucune ecriture vector_db ne sera effectuee sur ce run.")
            print("!" * 70 + "\n")
    else:
        print("\nMode configure: PHASE 1 ONLY (RUN_PHASE2=0)")

    succes = 0
    echecs = 0
    audit_entries = []
    tous_les_chunks_a_exporter = []

    print("\n" + "=" * 70)
    print("DEMARRAGE DE L'INGESTION")
    print("=" * 70)

    for filepath, type_doc in tqdm(fichiers_a_traiter, desc="Traitement", unit="doc"):
        try:
            indexed, extraction_method, extraction_info, raw_len, cleaned_len, status, doc_chunks = ingerer_document(
                filepath,
                type_doc,
                model,
                collection,
            )

            if doc_chunks:
                tous_les_chunks_a_exporter.extend(doc_chunks)

            ajouter_entree_audit(
                audit_entries,
                filepath,
                type_doc,
                extraction_method,
                extraction_info,
                raw_len,
                cleaned_len,
                status,
                indexed,
            )

            if indexed:
                succes += 1
            else:
                echecs += 1
        except Exception as e:
            print(f"Erreur avec {os.path.basename(filepath)}: {e}")
            ajouter_entree_audit(
                audit_entries,
                filepath,
                type_doc,
                "pipeline_exception",
                info_extraction(),
                0,
                0,
                "error",
                False,
            )
            echecs += 1

    print("\n" + "=" * 70)
    print("INGESTION TERMINEE")
    print("=" * 70)
    print(f"Succes (indexes): {succes}")
    print(f"Echecs (non indexes): {echecs}")
    if collection is not None:
        print(f"Total chunks ChromaDB: {collection.count()}")
    if collection is None:
        print("Total chunks en base: N/A (phase 2 non executee)")

    ecrire_rapport_audit(audit_entries)
    generer_rapport_couverture(audit_entries)
    
    if tous_les_chunks_a_exporter:
        exporter_chunks_vers_csv(tous_les_chunks_a_exporter)

    print("=" * 70 + "\n")


if __name__ == "__main__":
    ingerer_tous_documents()
